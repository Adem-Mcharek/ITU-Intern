"""
Meeting Notes Generator for ITU/UN Meetings
Generates professional .docx meeting notes with consistent formatting
based on transcript_speakers.txt files.
"""
import os
import json
from pathlib import Path
from datetime import datetime
from typing import Optional, Dict, Any
from flask import current_app

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("Warning: google-generativeai not available. Meeting notes generation will be disabled.")

try:
    from openai import AzureOpenAI
    AZURE_OPENAI_AVAILABLE = True
except ImportError:
    AZURE_OPENAI_AVAILABLE = False
    AzureOpenAI = None
    print("Warning: openai not available. Meeting notes generation will be disabled.")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not available. Meeting notes generation will be disabled.")

# Concise meeting notes prompt for ITU style - compact format
MEETING_NOTES_PROMPT = """
You are an ITU intern creating concise, professional meeting notes similar to UN/ITU diplomatic style.

CRITICAL: Keep each section SHORT and CONCISE.

Use this structure (make each section 1-2 short paragraphs or bullets ONLY):

**MEETING OVERVIEW**
1-2 sentences max: Purpose and key participants

**KEY DISCUSSIONS**
2-3 bullet points max. One sentence each.

**POSITIONS & RECOMMENDATIONS**  
2-3 bullet points max. Focus on consensus/disagreement only.

**DECISIONS & ACTION ITEMS**
• 2-3 decisions maximum
• 2-3 action items with responsible parties
• Timelines if any

**TECHNICAL MATTERS** (only if significant technical content)
1-2 bullets maximum

**CAPACITY BUILDING** (only if discussed)
1-2 bullets maximum

STYLE REQUIREMENTS:
- Use formal UN/ITU language but keep concise
- Third person: "[Representative] stated..." not "I stated..."
- Speaker attribution: "[Name, Organization] noted that..."
- Use ONLY bullet points in each section - NO long paragraphs
- Focus on substance, not process details
- Highlight key decisions and commitments

Generate concise meeting notes from this transcript:

"""

def setup_gemini_api() -> Optional[Any]:
    """Initialize Gemini API with configured key"""
    if not GEMINI_AVAILABLE:
        return None
    
    # Try to get API key from Flask config or environment
    api_key = None
    try:
        from flask import current_app
        api_key = current_app.config.get('GEMINI_API_KEY')
    except RuntimeError:
        # No Flask app context, try environment directly
        api_key = os.environ.get('GEMINI_API_KEY')
    
    if api_key:
        genai.configure(api_key=api_key)
        return genai.GenerativeModel("gemini-2.5-flash-lite-preview-06-17")
    return None


def setup_azure_openai_client():
    """Initialize Azure OpenAI client if configured"""
    if not AZURE_OPENAI_AVAILABLE:
        return None
    
    # Try to get config from Flask app or environment
    try:
        api_key = current_app.config.get('AZURE_OPENAI_API_KEY')
        endpoint = current_app.config.get('AZURE_OPENAI_ENDPOINT')
        api_version = current_app.config.get('AZURE_OPENAI_API_VERSION', '2024-12-01-preview')
        deployment = current_app.config.get('AZURE_OPENAI_DEPLOYMENT_NAME', 'GPT-4')
    except RuntimeError:
        api_key = os.environ.get('AZURE_OPENAI_API_KEY')
        endpoint = os.environ.get('AZURE_OPENAI_ENDPOINT')
        api_version = os.environ.get('AZURE_OPENAI_API_VERSION', '2024-12-01-preview')
        deployment = os.environ.get('AZURE_OPENAI_DEPLOYMENT_NAME', 'GPT-4')
    
    if not api_key or not endpoint:
        return None
    
    try:
        client = AzureOpenAI(
            api_key=api_key,
            api_version=api_version,
            azure_endpoint=endpoint
        )
        return client, deployment
    except Exception as e:
        print(f"Error initializing Azure OpenAI client: {e}")
        return None

def extract_meeting_metadata(speakers_file_path: Path, meeting_title: str) -> Dict[str, Any]:
    """Extract metadata from meeting content for document header"""
    try:
        with open(speakers_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract basic information
        speakers = []
        organizations = set()
        
        lines = content.split('\n')
        for line in lines:
            if line.strip().startswith('[') and line.strip().endswith(']'):
                # Parse speaker info
                speaker_info = line.strip()[1:-1]  # Remove brackets
                if ',' in speaker_info:
                    parts = speaker_info.split(',', 1)
                    speaker_name = parts[0].strip()
                    org = parts[1].strip()
                    speakers.append(speaker_name)
                    organizations.add(org)
                else:
                    speakers.append(speaker_info.strip())
        
        return {
            'title': meeting_title,
            'date': datetime.now().strftime('%B %d, %Y'),
            'total_speakers': len(set(speakers)),
            'organizations': list(organizations),
            'content_length': len(content)
        }
    
    except Exception as e:
        print(f"Error extracting meeting metadata: {e}")
        return {
            'title': meeting_title,
            'date': datetime.now().strftime('%B %d, %Y'),
            'total_speakers': 0,
            'organizations': [],
            'content_length': 0
        }


def extract_reference_format(meeting_dir: Path) -> Optional[str]:
    """Extract reference format from meeting notes reference document if it exists"""
    if not meeting_dir:
        return None
    
    # Check for reference documents (DOCX first, then PDF)
    reference_docx = meeting_dir / "meeting_notes_reference.docx"
    reference_pdf = meeting_dir / "meeting_notes_reference.pdf"
    
    if reference_docx.exists():
        return _extract_text_from_docx(reference_docx)
    elif reference_pdf.exists():
        return _extract_text_from_pdf(reference_pdf)
    
    return None


def _extract_text_from_docx(file_path: Path) -> Optional[str]:
    """Extract text from DOCX file"""
    try:
        from docx import Document
        doc = Document(str(file_path))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text if text.strip() else None
    except Exception as e:
        print(f"  Warning: Could not read reference DOCX: {e}")
        return None


def _extract_text_from_pdf(file_path: Path) -> Optional[str]:
    """Extract text from PDF file"""
    try:
        try:
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text if text.strip() else None
        except ImportError:
            # Try alternative PDF library
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
                return text if text.strip() else None
            except ImportError:
                print(f"  Warning: PDF libraries not available (install PyPDF2 or pdfplumber)")
                return None
    except Exception as e:
        print(f"  Warning: Could not read reference PDF: {e}")
        return None


def generate_meeting_notes_content(transcript_content: str, meeting_dir: Path = None) -> Optional[str]:
    """Generate professional meeting notes - tries Azure GPT-4 first, falls back to Gemini
    
    Args:
        transcript_content: The meeting transcript text
        meeting_dir: Optional meeting directory to check for reference format
    """
    if not transcript_content.strip():
        return None
    
    # Check for reference format if meeting_dir provided
    reference_format = None
    if meeting_dir:
        reference_format = extract_reference_format(meeting_dir)
        if reference_format:
            print("  ✓ Reference format found, using as template")
    
    # Prepare the full prompt
    if reference_format:
        # Use reference-based prompt
        full_prompt = f"""You are an ITU intern creating meeting notes.

CRITICAL: Follow this exact format from the reference example:

{reference_format}

Now generate meeting notes in exactly this format for the following transcript:

{transcript_content}

Generate meeting notes following the reference format strictly:"""
    else:
        # Use default prompt
        full_prompt = MEETING_NOTES_PROMPT + "\n\n" + transcript_content + "\n\nGenerate comprehensive meeting notes:"
    
    # Try Azure GPT-4 first (quick, minimal retries)
    print("Attempting Azure GPT-4...")
    result = _try_azure_openai(full_prompt)
    if result:
        return result
    
    # Fall back to Gemini if Azure fails
    print("Azure failed, trying Gemini fallback...")
    result = _try_gemini(full_prompt)
    if result:
        return result
    
    # Both failed
    print("Error: Both Azure GPT-4 and Gemini failed to generate meeting notes")
    return None


def _try_azure_openai(full_prompt: str) -> Optional[str]:
    """Try Azure OpenAI with quick retries (no long waits)"""
    client_info = setup_azure_openai_client()
    if not client_info:
        print("  Azure OpenAI not available, skipping...")
        return None
    
    client, deployment = client_info
    max_retries = 2  # Quick retries only
    
    for attempt in range(max_retries):
        try:
            print(f"  Azure attempt {attempt + 1}/{max_retries}...")
            
            response = client.chat.completions.create(
                model=deployment,
                messages=[
                    {"role": "system", "content": "You are an ITU intern creating professional meeting notes in UN/ITU diplomatic style."},
                    {"role": "user", "content": full_prompt}
                ],
                temperature=0.5,
                max_tokens=2000
            )
            
            notes_content = response.choices[0].message.content.strip()
            
            if len(notes_content) < 100:
                print(f"  Notes too short, retrying...")
                continue
            
            print(f"  ✓ Azure generated {len(notes_content)} characters")
            return notes_content
            
        except Exception as e:
            error_str = str(e)
            is_rate_limit = "429" in error_str or "RateLimitReached" in error_str
            
            if is_rate_limit:
                print(f"  ✗ Azure rate limited, will try Gemini")
                return None  # Exit to fallback
            else:
                print(f"  ✗ Azure error: {type(e).__name__}")
                if attempt < max_retries - 1:
                    import time
                    time.sleep(2)  # Quick 2 second wait
    
    return None


def _try_gemini(full_prompt: str) -> Optional[str]:
    """Fall back to Gemini API if Azure fails"""
    if not GEMINI_AVAILABLE:
        print("  Gemini not available")
        return None
    
    try:
        model = setup_gemini_api()
        if not model:
            print("  Gemini API not configured")
            return None
        
        print("  Trying Gemini API...")
        response = model.generate_content(full_prompt)
        
        notes_content = response.text.strip()
        
        if len(notes_content) < 100:
            print(f"  Gemini: Notes too short ({len(notes_content)} chars)")
            return None
        
        print(f"  ✓ Gemini generated {len(notes_content)} characters")
        return notes_content
        
    except Exception as e:
        print(f"  ✗ Gemini error: {type(e).__name__}: {e}")
        return None

def create_formatted_document(notes_content: str, metadata: Dict[str, Any]) -> Optional[Document]:
    """Create a professionally formatted Word document with the meeting notes"""
    if not DOCX_AVAILABLE:
        print("Warning: python-docx not available. Cannot create formatted document.")
        return None
    
    try:
        # Create new document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add ITU header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "International Telecommunication Union"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Document title
        title = doc.add_heading('MEETING NOTES', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Meeting title
        meeting_title = doc.add_heading(metadata['title'], 1)
        meeting_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add document info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(f"Date: {metadata['date']}")
        info_run.italic = True
        
        # Add separator line
        doc.add_paragraph("_" * 80).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add meeting overview if organizations present
        if metadata.get('organizations'):
            overview_para = doc.add_paragraph()
            overview_para.add_run("Participating Organizations: ").bold = True
            overview_para.add_run(", ".join(metadata['organizations'][:8]))  # Limit to first 8
            if len(metadata['organizations']) > 8:
                overview_para.add_run(" and others")
        
        # Add blank line
        doc.add_paragraph()
        
        # Process and add the generated content
        _add_formatted_content(doc, notes_content)
        
        # Add footer with generation info
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = f"Generated on {datetime.now().strftime('%B %d, %Y at %H:%M UTC')} | ITU INTERN (AI generated)"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return doc
        
    except Exception as e:
        print(f"Error creating formatted document: {e}")
        return None

def _add_formatted_content(doc: Document, content: str):
    """Add formatted content to the document with proper styling"""
    lines = content.split('\n')
    current_section = None
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check if line is a section header (starts with ** and ends with **)
        if line.startswith('**') and line.endswith('**'):
            section_title = line[2:-2].strip()
            heading = doc.add_heading(section_title, 2)
            heading.space_before = Pt(12)
            heading.space_after = Pt(6)
            current_section = section_title
            
        # Check if line is a subsection (starts with single bullet or dash)
        elif line.startswith('•') or line.startswith('-') or line.startswith('◦'):
            para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            para.space_after = Pt(3)
            
        # Check if line appears to be a speaker/participant identifier
        elif line.startswith('[') and ']' in line:
            para = doc.add_paragraph()
            speaker_run = para.add_run(line)
            speaker_run.bold = True
            speaker_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
            para.space_before = Pt(6)
            para.space_after = Pt(3)
            
        # Regular paragraph
        else:
            para = doc.add_paragraph(line)
            para.space_after = Pt(6)
            
            # Apply special formatting for certain keywords
            for run in para.runs:
                text = run.text.lower()
                if any(keyword in text for keyword in ['recommendation', 'decision', 'resolution', 'action item']):
                    run.bold = True

def save_meeting_notes_document(doc: Document, meeting_dir: Path, meeting_title: str) -> Optional[Path]:
    """Save the meeting notes document to the meeting directory"""
    try:
        # Create safe filename
        safe_title = "".join(c for c in meeting_title if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '_')[:50]  # Limit length
        
        filename = f"Meeting_Notes_{safe_title}_{datetime.now().strftime('%Y%m%d')}.docx"
        notes_path = meeting_dir / filename
        
        # Save document
        doc.save(str(notes_path))
        print(f"Meeting notes saved to: {notes_path}")
        
        return notes_path
        
    except Exception as e:
        print(f"Error saving meeting notes document: {e}")
        return None

def create_meeting_notes(meeting_id: int, speakers_file_path: Path, meeting_title: str) -> Optional[Path]:
    """
    Main function to create professional meeting notes
    
    Args:
        meeting_id: Database ID of the meeting
        speakers_file_path: Path to the transcript_speakers.txt file
        meeting_title: Title of the meeting
        
    Returns:
        Path to generated .docx file or None if failed
    """
    print(f"\n=== Generating Meeting Notes for Meeting {meeting_id} ===")
    
    # Check dependencies
    if not GEMINI_AVAILABLE or not DOCX_AVAILABLE:
        print("Error: Required dependencies not available (google-generativeai and/or python-docx)")
        return None
    
    # Check if file exists
    if not speakers_file_path.exists():
        print(f"Error: Transcript speakers file not found: {speakers_file_path}")
        return None
    
    # Extract metadata
    print("Step 1: Extracting meeting metadata...")
    metadata = extract_meeting_metadata(speakers_file_path, meeting_title)
    
    # Extract transcript content (reuse from summarizer)
    print("Step 2: Extracting transcript content...")
    from app.meeting_summarizer import extract_transcript_content
    transcript_content = extract_transcript_content(speakers_file_path)
    
    if not transcript_content:
        print("Error: No content extracted from transcript file")
        return None
    
    print(f"Extracted {len(transcript_content)} characters of transcript content")
    
    # Generate meeting notes content
    print("Step 3: Generating professional meeting notes with Azure GPT-4...")
    meeting_dir = speakers_file_path.parent
    notes_content = generate_meeting_notes_content(transcript_content, meeting_dir)
    
    if not notes_content:
        print("Error: Failed to generate meeting notes content")
        return None
    
    # Create formatted document
    print("Step 4: Creating professionally formatted Word document...")
    doc = create_formatted_document(notes_content, metadata)
    
    if not doc:
        print("Error: Failed to create formatted document")
        return None
    
    # Save document
    print("Step 5: Saving meeting notes document...")
    meeting_dir = speakers_file_path.parent
    notes_path = save_meeting_notes_document(doc, meeting_dir, meeting_title)
    
    if notes_path:
        print(f"✅ Successfully generated meeting notes: {notes_path}")
        return notes_path
    else:
        print("❌ Failed to save meeting notes document")
        return None

def save_notes_path_to_database(meeting_id: int, notes_path: Path) -> bool:
    """
    Save the meeting notes file path to the database
    
    Args:
        meeting_id: Database ID of the meeting
        notes_path: Path to the generated notes file
        
    Returns:
        True if saved successfully, False otherwise
    """
    try:
        from app import db
        from app.models import Meeting
        
        # Get meeting record
        meeting = Meeting.query.get(meeting_id)
        if not meeting:
            print(f"Error: Meeting {meeting_id} not found in database")
            return False
        
        # Update notes path field (we need to add this to the model)
        # For now, we'll store it relative to the uploads directory
        relative_path = notes_path.relative_to(notes_path.parent.parent)
        meeting.notes_path = str(relative_path)
        db.session.commit()
        
        print(f"✅ Saved meeting notes path to database for meeting {meeting_id}")
        return True
        
    except Exception as e:
        print(f"Error saving notes path to database: {e}")
        return False

def process_meeting_notes(meeting_id: int, meeting_dir: Path, meeting_title: str) -> bool:
    """
    Complete workflow to process meeting notes
    
    Args:
        meeting_id: Database ID of the meeting
        meeting_dir: Directory containing meeting files
        meeting_title: Title of the meeting
        
    Returns:
        True if processed successfully, False otherwise
    """
    try:
        # Locate speakers file
        speakers_file = meeting_dir / 'transcript_speakers.txt'
        
        # Generate notes
        notes_path = create_meeting_notes(meeting_id, speakers_file, meeting_title)
        
        if notes_path:
            # Save path to database
            return save_notes_path_to_database(meeting_id, notes_path)
        
        return False
        
    except Exception as e:
        print(f"Error in meeting notes workflow: {e}")
        return False